"""Convert Markdown to a structured Word document (.docx) — separate layout from PDF."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from typing import Any

from bs4 import NavigableString, Tag


def _utc_generated_line() -> str:
    now = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")
    return f"Generated: {now} · Cashflow Fraud Monitor"


def _add_cover_block(
    doc: Any,
    *,
    document_title: str,
    document_kind: str,
    classification: str,
    generated_line: str,
) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    p0 = doc.add_paragraph()
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(classification)
    r0.bold = True
    r0.font.size = Pt(8)
    r0.font.color.rgb = RGBColor(0x1A, 0x27, 0x44)

    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(document_kind)
    r1.bold = True
    r1.font.size = Pt(10)

    try:
        t = doc.add_paragraph(document_title, style="Title")
    except KeyError:
        t = doc.add_paragraph()
        tr = t.add_run(document_title)
        tr.bold = True
        tr.font.size = Pt(20)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    try:
        p2 = doc.add_paragraph(generated_line, style="Caption")
    except KeyError:
        p2 = doc.add_paragraph(generated_line)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    try:
        h = doc.add_paragraph("Narrative & findings (structured sections)", style="Heading 2")
    except KeyError:
        h = doc.add_paragraph()
        hr = h.add_run("Narrative & findings (structured sections)")
        hr.bold = True
    doc.add_paragraph()


def _add_inline_runs(paragraph: Any, node: Tag) -> None:
    from docx.shared import Pt, RGBColor

    for ch in node.children:
        if isinstance(ch, NavigableString):
            s = str(ch)
            if s:
                paragraph.add_run(s)
            continue
        if not isinstance(ch, Tag):
            continue
        name = (ch.name or "").lower()
        if name in ("strong", "b"):
            r = paragraph.add_run(ch.get_text())
            r.bold = True
        elif name in ("em", "i"):
            r = paragraph.add_run(ch.get_text())
            r.italic = True
        elif name == "code":
            r = paragraph.add_run(ch.get_text())
            r.font.name = "Courier New"
            r.font.size = Pt(10)
        elif name == "br":
            paragraph.add_run("\n")
        elif name == "a":
            r = paragraph.add_run(ch.get_text())
            r.font.color.rgb = RGBColor(0x05, 0x44, 0xAD)
            r.underline = True
        elif name in ("span", "mark", "del", "s"):
            _add_inline_runs(paragraph, ch)
        else:
            paragraph.add_run(ch.get_text())


def _add_table(doc: Any, table_elem: Tag) -> None:
    rows = table_elem.find_all("tr")
    if not rows:
        return
    col_counts = [len(r.find_all(["th", "td"])) for r in rows]
    ncols = max(col_counts) if col_counts else 0
    if ncols == 0:
        return
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    try:
        tbl.style = "Table Grid"
    except KeyError:
        pass
    for ri, tr in enumerate(rows):
        cells = tr.find_all(["th", "td"])
        for ci in range(ncols):
            cell = tbl.rows[ri].cells[ci]
            if ci < len(cells):
                td = cells[ci]
                cell.text = td.get_text("\n", strip=True)
            else:
                cell.text = ""


def _walk_blocks(doc: Any, root: Tag) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    for child in getattr(root, "children", ()) or ():
        if isinstance(child, NavigableString):
            t = str(child).strip()
            if t:
                doc.add_paragraph().add_run(t)
            continue
        if not isinstance(child, Tag):
            continue
        name = (child.name or "").lower()
        if name in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = min(max(int(name[1]), 1), 9)
            doc.add_heading(child.get_text(strip=True), level=level)
        elif name == "p":
            p = doc.add_paragraph()
            _add_inline_runs(p, child)
        elif name == "blockquote":
            p = doc.add_paragraph()
            _add_inline_runs(p, child)
            p.paragraph_format.left_indent = Inches(0.22)
            for r in p.runs:
                r.italic = True
        elif name == "pre":
            raw = child.get_text()
            for line in (raw.splitlines() or [""]):
                p = doc.add_paragraph()
                rr = p.add_run(line)
                rr.font.name = "Courier New"
                rr.font.size = Pt(9)
        elif name == "hr":
            p = doc.add_paragraph("—" * 32)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif name == "ul":
            for li in child.find_all("li", recursive=False):
                txt = li.get_text(" ", strip=True)
                if txt:
                    try:
                        doc.add_paragraph(txt, style="List Bullet")
                    except KeyError:
                        p = doc.add_paragraph()
                        p.add_run("• " + txt)
        elif name == "ol":
            for li in child.find_all("li", recursive=False):
                txt = li.get_text(" ", strip=True)
                if txt:
                    try:
                        doc.add_paragraph(txt, style="List Number")
                    except KeyError:
                        doc.add_paragraph(txt)
        elif name == "table":
            _add_table(doc, child)
        elif name in ("div", "section", "article", "tbody", "thead", "tfoot"):
            _walk_blocks(doc, child)
        elif name == "br":
            doc.add_paragraph()
        else:
            txt = child.get_text(" ", strip=True)
            if txt:
                doc.add_paragraph().add_run(txt)


def try_markdown_to_docx(
    markdown_text: str,
    *,
    document_title: str = "Financial crime narrative",
    document_kind: str = "Internal compliance worksheet",
    classification: str = "CONFIDENTIAL — Internal use only",
) -> tuple[bytes | None, str | None]:
    """
    Build a **Word** document from the same Markdown as the PDF path: cover sheet + structured body.

    This is a distinct deliverable (headings, lists, tables) rather than a PDF typography skin on HTML.
    """
    text = (markdown_text or "").strip()
    if not text:
        return None, "empty_markdown"

    try:
        import markdown
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        return None, f"missing_dependency:{e}"

    try:
        raw_html = markdown.markdown(text, extensions=["extra", "nl2br"])
    except Exception as e:  # noqa: BLE001
        return None, f"markdown_html:{type(e).__name__}:{e}"

    from bs4 import BeautifulSoup

    soup = BeautifulSoup(f"<div id='md-root'>{raw_html}</div>", "html.parser")
    root = soup.find("div", id="md-root")
    if root is None:
        return None, "html_parse_failed"

    doc = Document()
    _add_cover_block(
        doc,
        document_title=document_title,
        document_kind=document_kind,
        classification=classification,
        generated_line=_utc_generated_line(),
    )

    try:
        _walk_blocks(doc, root)
    except Exception as e:  # noqa: BLE001
        return None, f"docx_build:{type(e).__name__}:{e}"

    foot = doc.add_paragraph()
    fr = foot.add_run(
        "Disclaimer: Internal compliance worksheet generated from application data and, where applicable, "
        "model-assisted narrative. Not a filed Suspicious Activity Report (SAR). Verify all facts against "
        "source systems before any regulatory submission."
    )
    fr.font.size = Pt(8)

    buf = BytesIO()
    try:
        doc.save(buf)
    except Exception as e:  # noqa: BLE001
        return None, f"docx_save:{type(e).__name__}:{e}"

    out = buf.getvalue()
    if not out:
        return None, "empty_docx_output"
    return out, None
